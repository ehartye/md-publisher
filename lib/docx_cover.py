"""Per-theme cover page construction.

The PDF cover (cover.css) carries family-specific signals:
  - atlas:    italic Newsreader title, 1.2in red rule, uppercase Sora meta,
              attribution at bottom in small uppercase Sora
  - phosphor: lowercase mono title with '> ' accent prefix, '// ' subtitle
              prefix, '[ tag ]' bracketed meta items, mono attribution
  - arcade:   bold display title with hard 4px offset shadow in accent2,
              gradient rule, uppercase Bungee subtitle in cyan, neon meta

We translate each signal to the closest OOXML idiom we can hit:
  - hard offset shadow -> w:shadow run effect
  - gradient rule -> custom paragraph with thick bottom border in accent2
                     (gradient fills are not portable across Word versions)
  - prefixes ('> ', '// ', '[ ]') -> literal text in matching colors
  - 1.2in rule -> empty paragraph with sized bottom border, fixed-width
                  via a single tab whose width matches the rule

Cover lives in section 1 of the document; section 2 starts the body. The
cover section has the theme's bg color (already set globally by
styles.set_page_background) and a section break-after-page.

Per-family defaults can be overridden by the user theme's
``spec.json -> docx`` block (see spec). Currently honored:
  - ``accent_rule_thickness_pt``  (float; default 2.0)
  - ``cover_subtitle_caps``       (bool; default True for atlas/arcade)
"""

from __future__ import annotations

from docx.document import Document as _Doc
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from lxml import etree

from .docx_helpers import _w, hex_to_rgb
from .theme_loader import ThemeSelection


def _add_centered_paragraph(doc: _Doc, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_rule_paragraph(doc: _Doc, color_hex: str, *, width_in: float = 1.2,
                        thickness_pt: float = 1.5) -> None:
    """A short horizontal rule, theme-accent colored, centered.

    Implemented as a single run containing a unicode box-drawing character
    sized to approximate the rule width — Word's paragraph borders span
    the full text column width and can't be confined to N inches without
    section margin tricks. The U+2580 / U+2581 row of repeated characters
    gives a clean filled bar.
    """
    p = _add_centered_paragraph(doc)
    # Each box character is ~0.06" wide at the chosen pt — calibrate count.
    char = "▀"
    char_count = max(8, int(width_in / 0.06))
    run = p.add_run(char * char_count)
    run.font.size = Pt(thickness_pt * 4)  # vertical bar thickness
    run.font.color.rgb = RGBColor(*hex_to_rgb(color_hex))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(16)


def _add_gradient_rule_paragraph(doc: _Doc, color1: str, color2: str,
                                 *, width_in: float = 2.0,
                                 thickness_pt: float = 4.0) -> None:
    """Arcade's gradient rule — approximated as half-1 / half-2."""
    p = _add_centered_paragraph(doc)
    char = "▀"
    char_count = max(12, int(width_in / 0.06))
    half = char_count // 2
    r1 = p.add_run(char * half)
    r1.font.size = Pt(thickness_pt * 4)
    r1.font.color.rgb = RGBColor(*hex_to_rgb(color1))
    r2 = p.add_run(char * (char_count - half))
    r2.font.size = Pt(thickness_pt * 4)
    r2.font.color.rgb = RGBColor(*hex_to_rgb(color2))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(16)


def _add_vertical_space(doc: _Doc, pt: float) -> None:
    """Insert an empty paragraph of approximately N points height."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pt / 2)
    p.paragraph_format.space_after = Pt(pt / 2)


def _add_run_with_color(p, text: str, color_hex: str) -> None:
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*hex_to_rgb(color_hex))


def _set_run_shadow(run) -> None:
    """Apply Word's <w:shadow/> run effect — a hard offset shadow.

    Word renders this as a 1pt offset gray shadow; not the bold 4px-offset
    accent2 shadow the arcade PDF has, but it's the closest portable
    approximation in OOXML. (True multi-color text shadows require WordArt
    or text effects that don't round-trip across viewers.)
    """
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(_w("shadow")):
        rPr.remove(old)
    shadow = etree.SubElement(rPr, _w("shadow"))
    shadow.set(_w("val"), "1")


def add_cover_page(doc: _Doc, selection: ThemeSelection, *,
                   title: str, subtitle: str = "",
                   author: str = "", date: str = "",
                   attribution: str = "") -> None:
    """Emit the cover page contents into the current section of the doc.

    The doc must have its first section's margins set + page background
    applied before this is called. After the cover content, callers
    should invoke ``add_section_break(doc)`` so the body starts on page 2.

    Per-family defaults can be overridden via ``selection.spec_block['docx']``:
      - ``accent_rule_thickness_pt`` (float, default 2.0)
      - ``cover_subtitle_caps``       (bool, default True for atlas/arcade)
    """
    if selection.palette is None or selection.fonts is None:
        raise ValueError(
            "add_cover_page requires a themed ThemeSelection (palette and/or "
            "fonts is None — likely the 'default' theme was passed in; "
            "DOCX needs a real theme)."
        )

    docx_block = (selection.spec_block or {}).get("docx", {})
    accent_thickness = float(docx_block.get("accent_rule_thickness_pt", 2.0))
    subtitle_caps = bool(docx_block.get(
        "cover_subtitle_caps", selection.name in ("atlas", "arcade")
    ))

    # Top vertical breathing room — the cover.css uses 2.6in top margin.
    # We approximate with empty paragraphs since python-docx doesn't expose
    # vertical centering.
    _add_vertical_space(doc, 130)  # ~1.8 inch

    # ---- Title ----
    title_text = title
    title_prefix = ""
    if selection.name == "phosphor":
        title_text = title.lower()
        title_prefix = "> "
    p = _add_centered_paragraph(doc, style="MdpCoverTitle")
    if title_prefix:
        _add_run_with_color(p, title_prefix, selection.palette.accent)
    title_run = p.add_run(title_text)
    if selection.name == "arcade":
        # Closest OOXML approximation of the hard offset shadow signature.
        _set_run_shadow(title_run)

    # ---- Rule ----
    if selection.name == "arcade":
        _add_gradient_rule_paragraph(doc, selection.palette.accent_alt,
                                     selection.palette.accent,
                                     width_in=2.0, thickness_pt=accent_thickness)
    elif selection.name == "phosphor":
        _add_rule_paragraph(doc, selection.palette.ink_soft,
                            width_in=1.4, thickness_pt=0.5)
    else:
        _add_rule_paragraph(doc, selection.palette.accent,
                            width_in=1.2, thickness_pt=accent_thickness / 2)

    # ---- Subtitle ----
    if subtitle:
        sub_text = subtitle
        sub_prefix = ""
        if selection.name == "phosphor":
            sub_prefix = "// "
        elif subtitle_caps:
            sub_text = subtitle.upper()
        p = _add_centered_paragraph(doc, style="MdpCoverSubtitle")
        if sub_prefix:
            _add_run_with_color(p, sub_prefix, selection.palette.accent)
        p.add_run(sub_text)

    _add_vertical_space(doc, 16)

    # ---- Meta (author + date) ----
    meta_items = []
    if author:
        meta_items.append(author)
    if date:
        meta_items.append(date)

    for item in meta_items:
        p = _add_centered_paragraph(doc, style="MdpCoverMeta")
        if selection.name == "phosphor":
            # bracketed eyebrow style
            _add_run_with_color(p, "[ ", selection.palette.ink_soft)
            p.add_run(item)
            _add_run_with_color(p, " ]", selection.palette.ink_soft)
        else:
            p.add_run(item)

    # ---- Attribution at bottom ----
    if attribution:
        # Push to ~bottom of page
        _add_vertical_space(doc, 240)  # ~3.3 inch
        p = _add_centered_paragraph(doc, style="MdpCoverMeta")
        if selection.name == "arcade":
            # neon green attribution per cover.css
            _add_run_with_color(p, attribution.upper(),
                                selection.palette.accent_alt)
        elif selection.name == "atlas":
            # MdpCoverMeta carries <w:caps val="1"/> for atlas (set in
            # docx_styles._build_cover_styles), so the rendered text is
            # uppercase via Word's caps effect — leaving the underlying
            # text mixed-case so copy/paste + screen readers get the
            # original characters.
            p.add_run(attribution)
        else:
            p.add_run(attribution)


def add_section_break(doc: _Doc) -> None:
    """Start a new section after the cover so body content begins on page 2."""
    doc.add_section(WD_SECTION.NEW_PAGE)
