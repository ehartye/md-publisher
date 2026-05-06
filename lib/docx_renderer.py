"""Walk a markdown-it-py token stream and emit python-docx constructs.

Token types we handle (markdown-it-py emits these as a flat stream of
'open'/'close'/'inline' tokens):
  - heading_open / heading_close (h1 → 'Heading 1', etc.)
  - paragraph_open / paragraph_close
  - bullet_list_open / ordered_list_open / list_item_open (and their _close)
  - fence (```code``` blocks)
  - code_block (indented blocks — rare in this corpus)
  - blockquote_open / blockquote_close
  - hr
  - table_open / thead_open / tbody_open / tr_open / th_open / td_open
  - inline (children = text/em/strong/code_inline/link/image runs)
  - html_block — used by the productization pipeline for mermaid-PNG markers
    that ``lib.mermaid_processor`` injects upstream

We process the stream with a small index-based state machine because
markdown-it's tokens reference depth via .level, but the open/close pairs
are easy to bracket with a manual stack.

Mermaid input
-------------
Unlike the bake-off pydocx pipeline (which used a custom ``mdp_mermaid``
token type and a pre-rendered ``RenderedDiagram`` list), this port reads
the marker that ``lib.mermaid_processor`` emits into the markdown source:

    <img class='mermaid-png' src='file:///path/to/diagram.png' />

The renderer detects this marker inside ``html_block`` tokens, parses
the file URI, and embeds the PNG via ``Document.add_picture``. The dual
SVG+PNG embed path (``mermaid-dual``) lands in Task 3.2.
"""

from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import urlparse, unquote

from docx.document import Document as _Doc
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from lxml import etree
from markdown_it import MarkdownIt
from markdown_it.token import Token

from .docx_helpers import _w, hex_to_ooxml, hex_to_rgb
from .docx_syntax import LINE_BREAK, iter_highlighted_tokens
from .theme_loader import ThemeSelection


HEADING_MAP = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 4",
    "h6": "Heading 4",
}

# Default embedded mermaid PNG width in inches. 6.0" fits a 7.0"-content
# letter page (8.5" - 2*0.75" margins) comfortably with breathing room.
DEFAULT_MERMAID_WIDTH_IN = 6.0

# Marker patterns emitted by lib.mermaid_processor. We tolerate single OR
# double quotes around the src/data-svg attributes since the processor
# uses single quotes today but we shouldn't be that brittle.
_MERMAID_PNG_RE = re.compile(
    r"""<img\s+class=['"]mermaid-png['"]\s+src=['"]([^'"]+)['"]\s*/?>""",
    re.IGNORECASE,
)
_MERMAID_DUAL_RE = re.compile(
    r"""<img\s+class=['"]mermaid-dual['"]\s+data-svg=['"]([^'"]+)['"]\s*/?>""",
    re.IGNORECASE,
)


def _file_uri_to_path(uri: str) -> str:
    """Convert a 'file:///c:/path/to/x.png' URI to a local filesystem path.

    python-docx's add_picture wants a string path, not a URI. urlparse
    gives us the cross-platform decoding (handles %20 etc. via unquote).
    """
    parsed = urlparse(uri)
    if parsed.scheme not in ("file", ""):
        return uri  # not a local file URI; pass through and let docx error
    path = unquote(parsed.path)
    # Windows: urlparse('file:///c:/x') gives path '/c:/x'; strip leading /.
    if re.match(r"^/[a-zA-Z]:[\\/]", path):
        path = path[1:]
    return path


def render_inline(container, tokens: list[Token], selection: ThemeSelection,
                  base_style: str | None = None) -> None:
    """Emit child runs of an `inline` token into a paragraph (or cell para).

    container is anything with .add_run() — Paragraph or table-cell paragraph.
    We track em/strong/code state via a stack since markdown-it nests opens
    arbitrarily.
    """
    em = 0
    strong = 0
    link_href: list[str] = []  # stack
    for tok in tokens:
        t = tok.type
        if t == "text":
            text = tok.content
            if not text:
                continue
            run = container.add_run(text)
            if em:
                run.italic = True
            if strong:
                run.bold = True
            if link_href:
                # underline + accent color to signal a link (we don't
                # construct a real hyperlink relationship here; it's a
                # cosmetic link for the printed PDF-equivalent output)
                run.font.color.rgb = RGBColor(*hex_to_rgb(selection.palette.accent))
                run.underline = True
        elif t == "em_open":
            em += 1
        elif t == "em_close":
            em = max(0, em - 1)
        elif t == "strong_open":
            strong += 1
        elif t == "strong_close":
            strong = max(0, strong - 1)
        elif t == "code_inline":
            run = container.add_run(tok.content)
            run.font.name = selection.fonts.mono
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*hex_to_rgb(selection.palette.code_text))
            # apply <w:rFonts cs/eastAsia/..> too via the run's rPr (the
            # high-level python-docx API only sets ascii/hAnsi)
            rPr = run._element.get_or_add_rPr()
            for old in rPr.findall(_w("rFonts")):
                rPr.remove(old)
            rfonts = etree.SubElement(rPr, _w("rFonts"))
            rfonts.set(_w("ascii"), selection.fonts.mono)
            rfonts.set(_w("hAnsi"), selection.fonts.mono)
            rfonts.set(_w("cs"), selection.fonts.mono)
            rfonts.set(_w("eastAsia"), selection.fonts.mono)
            # subtle inline-code background
            for old in rPr.findall(_w("shd")):
                rPr.remove(old)
            shd = etree.SubElement(rPr, _w("shd"))
            shd.set(_w("val"), "clear")
            shd.set(_w("color"), "auto")
            shd.set(_w("fill"), hex_to_ooxml(selection.palette.code_bg))
        elif t == "link_open":
            link_href.append(tok.attrs.get("href", ""))
        elif t == "link_close":
            if link_href:
                link_href.pop()
        elif t == "softbreak":
            container.add_run(" ")
        elif t == "hardbreak":
            container.add_run().add_break()
        # image_inline / html_inline left unhandled — corpus doesn't use them


def _add_horizontal_rule(doc: _Doc, selection: ThemeSelection) -> None:
    """Emit a paragraph with a bottom border — Word's idiom for an HR."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, _w("pBdr"))
    bottom = etree.SubElement(pBdr, _w("bottom"))
    bottom.set(_w("val"), "single")
    bottom.set(_w("sz"), "8")
    bottom.set(_w("space"), "1")
    bottom.set(_w("color"), hex_to_ooxml(selection.palette.accent))


def _add_code_fence(doc: _Doc, content: str, selection: ThemeSelection,
                    lang_info: str = "") -> None:
    """Emit a code block as ONE paragraph with soft line breaks + highlighting.

    Previously we emitted one paragraph per source line so every line
    carried its own shading + accent rule. That stacks N rectangles and N
    border strokes — visually fragmented vs the WeasyPrint PDF, which
    renders a fenced block as one continuous shaded panel.

    Now we emit a single ``MdpCodeBlock`` paragraph and separate source
    lines with ``<w:br/>`` (Word soft line breaks). One paragraph means
    one continuous shaded rectangle + one continuous left border.

    Pygments tokenizes the fence by language hint; per-token color, bold,
    and italic come from the per-mode pygments.css palette.
    """
    p = doc.add_paragraph(style="MdpCodeBlock")
    lang = (lang_info or "").strip().split()[0] if lang_info else ""

    for text, color, bold, italic in iter_highlighted_tokens(content, lang, selection.mode):
        if text == LINE_BREAK:
            br_run = p.add_run()
            # Default <w:br/> with no w:type is a line break (soft break),
            # which is exactly what we want here.
            etree.SubElement(br_run._element, _w("br"))
            continue
        run = p.add_run(text)
        # Preserve leading/internal whitespace in the run's <w:t> nodes —
        # without xml:space="preserve" Word collapses leading runs of
        # spaces, breaking code indentation.
        for t_el in run._element.findall(_w("t")):
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        if color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(color))
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def _add_heading_with_rule(doc: _Doc, text: str, level: int,
                           selection: ThemeSelection) -> None:
    """Add a heading; Heading 1 also gets a thin accent rule below (atlas)."""
    style = HEADING_MAP[f"h{min(level, 6)}"]
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    if level == 1 and selection.name == "atlas":
        # Atlas signature: 1px accent rule under H1
        pPr = p._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, _w("pBdr"))
        bottom = etree.SubElement(pBdr, _w("bottom"))
        bottom.set(_w("val"), "single")
        bottom.set(_w("sz"), "6")
        bottom.set(_w("space"), "4")
        bottom.set(_w("color"), hex_to_ooxml(selection.palette.accent))


def _list_paragraph(doc: _Doc, text_tokens: list[Token],
                    selection: ThemeSelection,
                    ordered: bool, depth: int) -> None:
    """Emit a single list item.

    We use the built-in 'List Bullet' / 'List Number' styles which Word
    knows how to render with bullets/numbering. For nested lists we add
    indentation manually since Word's list numbering definitions are a
    rabbit hole disproportionate to the value here.
    """
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        # Some Word installs don't ship 'List Number' as a built-in style
        # name — fall back to Normal with a manual marker.
        p = doc.add_paragraph()
        marker = "1. " if ordered else "• "
        p.add_run(marker)
    if depth > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
    render_inline(p, text_tokens, selection)


def _shade_cell(cell: _Cell, fill_hex: str) -> None:
    """Apply a background fill to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(_w("shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, _w("shd"))
    shd.set(_w("val"), "clear")
    shd.set(_w("color"), "auto")
    shd.set(_w("fill"), hex_to_ooxml(fill_hex))


def _embed_mermaid_png(doc: _Doc, png_path: str,
                       *, width_in: float = DEFAULT_MERMAID_WIDTH_IN) -> None:
    """Embed a mermaid PNG as a centered, fixed-width inline picture."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Inches(width_in))


def _embed_missing_diagram_placeholder(doc: _Doc, png_path: str) -> None:
    """Emit a visible placeholder paragraph when a referenced diagram is gone.

    Reachable when (a) the build cache was cleaned between the mermaid
    preprocess and the docx render, (b) the cache lives on a removable
    drive that became unmounted, or (c) someone hand-edits the markdown
    to paste a stale marker. Without this fallback, add_picture would
    raise FileNotFoundError mid-render and corrupt the output.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[mermaid diagram missing: {png_path}]")
    run.italic = True


def _try_handle_mermaid_html(doc: _Doc, html_content: str) -> bool:
    """If html_content contains a mermaid-png marker, embed it; return True.

    Returns False for any other html_block content (caller drops silently).
    The dual-embed marker (mermaid-dual) is recognized but deferred to
    Task #14 — for now we extract the SVG path and embed via the PNG path
    by rasterizing on demand (lands in 14, currently a no-op skip).

    Failure mode: if the marker references a missing PNG file (cache wiped,
    unmounted drive, stale hand-edit), we emit a visible placeholder
    paragraph rather than crashing the whole render.
    """
    m = _MERMAID_PNG_RE.search(html_content)
    if m:
        png_path = _file_uri_to_path(m.group(1))
        if not Path(png_path).is_file():
            _embed_missing_diagram_placeholder(doc, png_path)
        else:
            _embed_mermaid_png(doc, png_path)
        return True
    # mermaid-dual: Task #14 will fan out to SVG+PNG embed. For Phase 2,
    # ignore — the PNG path is the supported one.
    if _MERMAID_DUAL_RE.search(html_content):
        return True  # consumed (deferred); don't fall through to "drop"
    return False


def render_tokens(doc: _Doc, selection: ThemeSelection,
                  tokens: list[Token]) -> None:
    """Walk the top-level token stream and emit Word constructs.

    The stream is roughly: paragraph_open, inline (with .children), paragraph_close
    repeated, with occasional fence/heading/list/blockquote/table/hr blocks.
    """
    i = 0
    n = len(tokens)
    list_depth = -1  # -1 = not in a list
    list_ordered_stack: list[bool] = []

    while i < n:
        tok = tokens[i]
        t = tok.type

        if t == "heading_open":
            level = int(tok.tag[1:])
            inline = tokens[i + 1]
            text = "".join(c.content for c in (inline.children or []) if c.type == "text")
            _add_heading_with_rule(doc, text, level, selection)
            i += 3  # heading_open, inline, heading_close
            continue

        if t == "paragraph_open":
            inline = tokens[i + 1]
            # In list items, paragraph wraps the item text; we let
            # _list_paragraph handle that by skipping paragraph_open in lists.
            if list_depth >= 0:
                _list_paragraph(doc, inline.children or [], selection,
                                list_ordered_stack[-1], list_depth)
            else:
                p = doc.add_paragraph()
                render_inline(p, inline.children or [], selection)
            i += 3
            continue

        if t == "fence":
            _add_code_fence(doc, tok.content, selection, lang_info=tok.info or "")
            i += 1
            continue

        if t == "code_block":
            # Indented (4-space) blocks have no language hint by definition.
            _add_code_fence(doc, tok.content, selection, lang_info="")
            i += 1
            continue

        if t == "hr":
            _add_horizontal_rule(doc, selection)
            i += 1
            continue

        if t == "blockquote_open":
            # find matching close, render contents as MdpBlockquote paragraphs
            depth = 1
            j = i + 1
            while j < n and depth:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                j += 1
            inner = tokens[i + 1:j - 1]
            k = 0
            while k < len(inner):
                if inner[k].type == "paragraph_open":
                    inline = inner[k + 1]
                    p = doc.add_paragraph(style="MdpBlockquote")
                    render_inline(p, inline.children or [], selection)
                    k += 3
                else:
                    k += 1
            i = j
            continue

        if t in ("bullet_list_open", "ordered_list_open"):
            list_depth += 1
            list_ordered_stack.append(t == "ordered_list_open")
            i += 1
            continue
        if t in ("bullet_list_close", "ordered_list_close"):
            list_depth -= 1
            if list_ordered_stack:
                list_ordered_stack.pop()
            i += 1
            continue
        if t in ("list_item_open", "list_item_close"):
            i += 1
            continue

        if t == "table_open":
            j = i + 1
            depth = 1
            while j < n and depth:
                if tokens[j].type == "table_open":
                    depth += 1
                elif tokens[j].type == "table_close":
                    depth -= 1
                j += 1
            _render_table(doc, selection, tokens[i + 1:j - 1])
            i = j
            continue

        if t == "html_block":
            # html_block is the carrier for mermaid-png markers from
            # lib.mermaid_processor. Anything else (or unrecognized) is
            # dropped silently.
            _try_handle_mermaid_html(doc, tok.content or "")
            i += 1
            continue

        # Unhandled token — skip silently. The corpus doesn't use any
        # token types we haven't covered, so this is defensive.
        i += 1


def _render_table(doc: _Doc, selection: ThemeSelection,
                  tokens: list[Token]) -> None:
    """Walk a table token subrange and emit a docx table with theme styling."""
    # First pass: collect header rows (tr inside thead) and body rows (tr in tbody)
    header_rows: list[list[list[Token]]] = []  # row -> cell -> inline.children
    body_rows: list[list[list[Token]]] = []
    in_thead = False
    in_tbody = False
    cur_row: list[list[Token]] | None = None
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "thead_open":
            in_thead = True
        elif tok.type == "thead_close":
            in_thead = False
        elif tok.type == "tbody_open":
            in_tbody = True
        elif tok.type == "tbody_close":
            in_tbody = False
        elif tok.type == "tr_open":
            cur_row = []
        elif tok.type == "tr_close":
            if cur_row is not None:
                (header_rows if in_thead else body_rows).append(cur_row)
                cur_row = None
        elif tok.type in ("th_open", "td_open"):
            inline = tokens[i + 1]
            children = inline.children or [] if inline.type == "inline" else []
            cur_row.append(children)
            i += 2
        i += 1

    if not header_rows and not body_rows:
        return
    n_cols = max(len(r) for r in header_rows + body_rows)
    n_rows = len(header_rows) + len(body_rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = True

    row_idx = 0
    for hrow in header_rows:
        for col_idx, cell_tokens in enumerate(hrow):
            cell = table.rows[row_idx].cells[col_idx]
            header_fill = (selection.palette.code_bg if selection.name != "arcade"
                           else selection.palette.accent_alt)
            _shade_cell(cell, header_fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # clear default paragraph
            p = cell.paragraphs[0]
            for run in list(p.runs):
                run.text = ""
            render_inline(p, cell_tokens, selection)
            for run in p.runs:
                run.bold = True
                if selection.name == "arcade":
                    arcade_header_color = ("#FFFFFF" if selection.mode == "dark"
                                           else selection.palette.ink)
                    run.font.color.rgb = RGBColor(*hex_to_rgb(arcade_header_color))
        row_idx += 1
    for brow in body_rows:
        for col_idx, cell_tokens in enumerate(brow):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # subtle stripe on alt rows
            if (row_idx - len(header_rows)) % 2 == 1:
                _shade_cell(cell, selection.palette.code_bg)
            p = cell.paragraphs[0]
            for run in list(p.runs):
                run.text = ""
            render_inline(p, cell_tokens, selection)
        row_idx += 1


def parse_markdown(md_text: str) -> list[Token]:
    """Parse markdown to tokens, with table + strikethrough enabled.

    Mermaid fences should already have been preprocessed upstream by
    ``lib.mermaid_processor``, which replaces them with
    ``<img class='mermaid-png' .../>`` markers (consumed in render_tokens
    via the html_block branch). We don't intercept ```mermaid here —
    if the upstream processor didn't run, the raw fence will appear
    as a regular code block.
    """
    md = MarkdownIt("commonmark").enable("table").enable("strikethrough")
    return md.parse(md_text)


def render_markdown_to_docx(doc: _Doc, selection: ThemeSelection,
                            md_text: str) -> None:
    """Public entry point — parse md_text and emit Word constructs into doc.

    Caller is responsible for:
      - Pre-running ``lib.mermaid_processor.MermaidProcessor.run()`` on
        the markdown lines if mermaid embedding is desired (PNG mode).
      - Applying theme styles via ``lib.docx_styles.apply_theme()`` first.
      - Adding cover page + TOC before this if desired.
      - Saving the doc afterward.
    """
    if selection.palette is None or selection.fonts is None:
        raise ValueError(
            "render_markdown_to_docx requires a themed ThemeSelection "
            "(palette and/or fonts is None — likely the 'default' theme "
            "was passed in; DOCX needs a real theme)."
        )

    tokens = parse_markdown(md_text)
    render_tokens(doc, selection, tokens)
