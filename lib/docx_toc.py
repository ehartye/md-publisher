"""Insert a Word native TOC field that auto-populates on first open.

A Word table-of-contents is a field instruction (`TOC \\o "1-3" \\h \\z \\u`)
wrapped in a w:fldSimple element. When the document is opened in Word, the
user gets a one-time "update fields" prompt; if they accept, Word walks the
heading styles and writes the TOC body. If they decline, the field shows a
'No table of contents entries found.' placeholder and the user can right-
click → Update Field later.

For the bake-off this is the right tradeoff: we get clickable, page-numbered,
auto-styled TOC entries with zero render-time work, and the build doesn't
have to walk the document twice to compute page numbers (which would be
wrong anyway since Word does its own pagination).
"""

from __future__ import annotations

from docx.document import Document as _Doc
from docx.shared import RGBColor
from lxml import etree

from .docx_helpers import _w, hex_to_rgb
from .theme_loader import ThemeSelection


def insert_toc(doc: _Doc, selection: ThemeSelection) -> None:
    """Add a TOC heading and an auto-updating TOC field."""
    # Heading for the TOC
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Contents")

    # The TOC field. We use the begin/separate/end form because some Word
    # versions don't refresh fldSimple TOCs reliably; the long form is
    # what Word itself emits when you Insert → Table of Contents.
    p = doc.add_paragraph()
    run = p.add_run()
    r_el = run._element

    # <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
    fldChar1 = etree.SubElement(r_el, _w("fldChar"))
    fldChar1.set(_w("fldCharType"), "begin")
    fldChar1.set(_w("dirty"), "true")

    # <w:r><w:instrText xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText></w:r>
    run2 = p.add_run()
    r_el2 = run2._element
    instr = etree.SubElement(r_el2, _w("instrText"))
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    # <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    run3 = p.add_run()
    r_el3 = run3._element
    fldChar2 = etree.SubElement(r_el3, _w("fldChar"))
    fldChar2.set(_w("fldCharType"), "separate")

    # Placeholder text shown until Word updates the field
    run4 = p.add_run("Right-click → Update Field to populate the table of contents.")
    run4.italic = True
    run4.font.color.rgb = RGBColor(*hex_to_rgb(selection.palette.ink_soft))

    # <w:r><w:fldChar w:fldCharType="end"/></w:r>
    run5 = p.add_run()
    r_el5 = run5._element
    fldChar3 = etree.SubElement(r_el5, _w("fldChar"))
    fldChar3.set(_w("fldCharType"), "end")

    # Page break after the TOC
    doc.add_page_break()
